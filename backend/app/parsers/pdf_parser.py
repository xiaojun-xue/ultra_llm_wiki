"""PDF and Word document parser."""

import re

from app.parsers.base import BaseParser, ParseResult, ParsedChunk


class PDFParser(BaseParser):
    def supported_extensions(self) -> set[str]:
        return {".pdf"}

    async def parse(self, data: bytes, filename: str) -> ParseResult:
        import pymupdf

        doc = pymupdf.open(stream=data, filetype="pdf")
        pages_text = []
        all_text_parts = []

        for page_num, page in enumerate(doc):
            text = page.get_text("text")
            if text.strip():
                pages_text.append((page_num + 1, text))
                all_text_parts.append(text)

        full_text = "\n\n".join(all_text_parts)
        doc.close()

        # If very little text extracted, the PDF is likely image-based
        if len(full_text.strip()) < 100 and len(list(doc)) > 0:
            # Re-open and try OCR-based extraction via image rendering
            full_text, pages_text = await self._ocr_fallback(data)

        # Build chunks per page (or merge small pages)
        chunks = self._build_chunks(pages_text, filename)

        # Try to get title from first page
        title = filename
        if pages_text:
            first_lines = pages_text[0][1].strip().split("\n")
            if first_lines and len(first_lines[0]) < 200:
                title = first_lines[0].strip() or filename

        refs = self._extract_references(full_text)

        return ParseResult(
            title=title,
            doc_type="document",
            content=full_text,
            chunks=chunks,
            metadata={"format": "pdf", "pages": len(pages_text)},
            references=refs,
        )

    async def _ocr_fallback(self, data: bytes) -> tuple[str, list[tuple[int, str]]]:
        """Fallback: render PDF pages as images and OCR them."""
        # OCR requires paddleocr or tesseract; return empty if not available
        try:
            import pymupdf

            doc = pymupdf.open(stream=data, filetype="pdf")
            pages_text = []
            for page_num, page in enumerate(doc):
                # Render page at 200 DPI
                pix = page.get_pixmap(dpi=200)
                img_data = pix.tobytes("png")
                text = await self._ocr_image(img_data)
                if text.strip():
                    pages_text.append((page_num + 1, text))
            doc.close()
            full_text = "\n\n".join(t for _, t in pages_text)
            return full_text, pages_text
        except Exception:
            return "", []

    async def _ocr_image(self, img_data: bytes) -> str:
        """OCR a single image. Override this with PaddleOCR/Tesseract."""
        # Placeholder: in production, use PaddleOCR for Chinese/English
        # from paddleocr import PaddleOCR
        # ocr = PaddleOCR(use_angle_cls=True, lang='ch')
        # result = ocr.ocr(img_data)
        return ""

    def _build_chunks(
        self, pages_text: list[tuple[int, str]], filename: str
    ) -> list[ParsedChunk]:
        """Merge small pages, split large ones."""
        chunks = []
        buffer = []
        buffer_len = 0
        buffer_start_page = 1

        for page_num, text in pages_text:
            text = text.strip()
            if not text:
                continue

            if buffer_len == 0:
                buffer_start_page = page_num

            buffer.append(text)
            buffer_len += len(text)

            # Flush if accumulated enough content
            if buffer_len >= 1500:
                chunks.append(ParsedChunk(
                    content="\n\n".join(buffer),
                    metadata={
                        "type": "pdf_pages",
                        "file": filename,
                        "start_page": buffer_start_page,
                        "end_page": page_num,
                    },
                ))
                buffer = []
                buffer_len = 0

        # Flush remaining
        if buffer:
            chunks.append(ParsedChunk(
                content="\n\n".join(buffer),
                metadata={
                    "type": "pdf_pages",
                    "file": filename,
                    "start_page": buffer_start_page,
                    "end_page": pages_text[-1][0] if pages_text else buffer_start_page,
                },
            ))

        return chunks

    def _extract_references(self, text: str) -> list[str]:
        """Extract filenames and references from PDF text."""
        refs = []
        # Look for filename-like patterns
        for m in re.finditer(r"[\w/\\]+\.\w{1,5}", text):
            candidate = m.group(0)
            if any(candidate.endswith(ext) for ext in (".c", ".h", ".py", ".java", ".sch", ".pdf")):
                refs.append(candidate)
        return list(set(refs))


class WordParser(BaseParser):
    def supported_extensions(self) -> set[str]:
        return {".docx", ".doc"}

    async def parse(self, data: bytes, filename: str) -> ParseResult:
        from io import BytesIO

        from docx import Document as DocxDocument

        doc = DocxDocument(BytesIO(data))

        paragraphs = []
        chunks = []
        current_section = []
        current_heading = ""

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            paragraphs.append(text)

            # Check if this is a heading
            if para.style and para.style.name.startswith("Heading"):
                # Flush previous section
                if current_section:
                    section_text = "\n\n".join(current_section)
                    if current_heading:
                        section_text = f"[{current_heading}]\n\n{section_text}"
                    chunks.append(ParsedChunk(
                        content=section_text,
                        metadata={"type": "section", "heading": current_heading},
                    ))
                    current_section = []
                current_heading = text
                current_section.append(text)
            else:
                current_section.append(text)

        # Flush last section
        if current_section:
            section_text = "\n\n".join(current_section)
            if current_heading:
                section_text = f"[{current_heading}]\n\n{section_text}"
            chunks.append(ParsedChunk(
                content=section_text,
                metadata={"type": "section", "heading": current_heading},
            ))

        full_text = "\n\n".join(paragraphs)

        # Title from first heading or filename
        title = filename
        for para in doc.paragraphs:
            if para.style and para.style.name.startswith("Heading"):
                title = para.text.strip()
                break

        return ParseResult(
            title=title,
            doc_type="document",
            content=full_text,
            chunks=chunks or [ParsedChunk(content=full_text, metadata={"type": "full"})],
            metadata={"format": "docx"},
        )
